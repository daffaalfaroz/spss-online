import json
import os
import io
import pandas as pd
import numpy as np
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth.models import User
from django.http import JsonResponse, HttpResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.contrib import messages

from .models import Dataset, AnalysisResult, ChartResult
from analysis.engine import AnalysisEngine
from analysis.charts import ChartGenerator


# ──────────────────────────────────────────────────────────────────────────────
# AUTH VIEWS
# ──────────────────────────────────────────────────────────────────────────────

def register_view(request):
    if request.user.is_authenticated:
        return redirect('editor')
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            messages.success(request, f'Selamat datang, {user.username}!')
            return redirect('editor')
    else:
        form = UserCreationForm()
    return render(request, 'auth/register.html', {'form': form})


def login_view(request):
    if request.user.is_authenticated:
        return redirect('editor')
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('editor')
    else:
        form = AuthenticationForm()
    return render(request, 'auth/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('login')


# ──────────────────────────────────────────────────────────────────────────────
# MAIN EDITOR
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def editor_view(request):
    datasets = Dataset.objects.filter(user=request.user).order_by('-created_at')[:10]
    return render(request, 'core/editor.html', {'datasets': datasets})


# ──────────────────────────────────────────────────────────────────────────────
# DATA API
# ──────────────────────────────────────────────────────────────────────────────

def load_df(dataset):
    """Load dataset as pandas DataFrame."""
    fpath = dataset.file.path
    ext = os.path.splitext(fpath)[1].lower()
    if ext in ['.xlsx', '.xls']:
        df = pd.read_excel(fpath)
    elif ext == '.sav':
        import pyreadstat
        df, meta = pyreadstat.read_sav(fpath)
    else:
        # Try multiple encodings
        df = None
        for enc in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(fpath, encoding=enc)
                break
            except Exception:
                continue
        if df is None:
            df = pd.read_csv(fpath, encoding='utf-8', errors='replace')
    # Ensure all column names are strings
    df.columns = [str(c) for c in df.columns]
    return df


def df_to_safe_json(df):
    """Convert DataFrame to JSON-safe list of lists, handling NaN/Inf correctly."""
    # pandas to_json handles NaN -> null natively
    return json.loads(df.to_json(orient='values', default_handler=str))


@login_required
def upload_dataset(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    file = request.FILES.get('file')
    if not file:
        return JsonResponse({'error': 'No file provided'}, status=400)

    ext = os.path.splitext(file.name)[1].lower()
    if ext not in ['.csv', '.xlsx', '.xls', '.sav']:
        return JsonResponse({'error': 'Format file tidak didukung. Gunakan CSV, Excel, atau SPSS (.sav)'}, status=400)

    try:
        dataset = Dataset(
            user=request.user,
            name=os.path.splitext(file.name)[0],
            original_filename=file.name,
            file_type=ext.lstrip('.'),
        )
        dataset.file.save(file.name, file)
        dataset.save()

        df = load_df(dataset)
        dataset.rows = len(df)
        dataset.columns = len(df.columns)

        col_info = []
        for col in df.columns:
            dtype = str(df[col].dtype)
            if 'int' in dtype or 'float' in dtype:
                col_type = 'numeric'
            elif 'datetime' in dtype:
                col_type = 'date'
            else:
                col_type = 'string'
            col_info.append({
                'name': col,
                'type': col_type,
                'dtype': dtype,
                'missing': int(df[col].isna().sum()),
            })
        dataset.set_column_info(col_info)
        dataset.save()

        # Return preview — use pandas JSON serializer to safely handle NaN -> null
        preview_df = df.head(100)
        preview_data = df_to_safe_json(preview_df)

        return JsonResponse({
            'success': True,
            'dataset_id': dataset.id,
            'name': dataset.name,
            'rows': dataset.rows,
            'columns': dataset.columns,
            'column_info': col_info,
            'preview_columns': list(df.columns),
            'preview_data': preview_data,
        })
    except Exception as e:
        return JsonResponse({'error': f'Error membaca file: {str(e)}'}, status=400)


@login_required
def get_dataset(request, dataset_id):
    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)
    df = load_df(dataset)
    data = df_to_safe_json(df)
    return JsonResponse({
        'success': True,
        'dataset_id': dataset.id,
        'name': dataset.name,
        'rows': dataset.rows,
        'columns': dataset.columns,
        'column_info': dataset.get_column_info(),
        'columns_list': list(df.columns),
        'data': data,
    })


@login_required
def list_datasets(request):
    datasets = Dataset.objects.filter(user=request.user).order_by('-created_at')
    result = []
    for ds in datasets:
        result.append({
            'id': ds.id,
            'name': ds.name,
            'filename': ds.original_filename,
            'rows': ds.rows,
            'columns': ds.columns,
            'created_at': ds.created_at.strftime('%d/%m/%Y %H:%M'),
        })
    return JsonResponse({'datasets': result})


@login_required
def delete_dataset(request, dataset_id):
    if request.method != 'DELETE':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)
    try:
        if dataset.file and os.path.exists(dataset.file.path):
            os.remove(dataset.file.path)
    except Exception:
        pass
    dataset.delete()
    return JsonResponse({'success': True})


@login_required
def save_dataset(request, dataset_id):
    """Save modified dataset (edited cells, added/removed rows & columns)."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)

    try:
        body = json.loads(request.body)
        columns = body.get('columns', [])
        data = body.get('data', [])

        if not columns:
            return JsonResponse({'error': 'Columns tidak boleh kosong'}, status=400)
        if not data:
            data = [[''] * len(columns)]

        df = pd.DataFrame(data, columns=columns)

        # Try to infer numeric types
        for col in df.columns:
            try:
                converted = pd.to_numeric(df[col], errors='coerce')
                if converted.notna().sum() > 0:
                    df[col] = converted
            except Exception:
                pass

        # Save back to file
        fpath = dataset.file.path
        ext = os.path.splitext(fpath)[1].lower()

        if ext in ['.xlsx', '.xls']:
            df.to_excel(fpath, index=False)
        elif ext == '.sav':
            # Convert sav → csv since pyreadstat write needs extra setup
            csv_path = os.path.splitext(fpath)[0] + '_edited.csv'
            df.to_csv(csv_path, index=False, encoding='utf-8-sig')
            from django.conf import settings as djsettings
            dataset.file.name = os.path.relpath(csv_path, djsettings.MEDIA_ROOT).replace('\\', '/')
        else:
            df.to_csv(fpath, index=False, encoding='utf-8-sig')

        # Rebuild column info
        dataset.rows = len(df)
        dataset.columns = len(df.columns)
        col_info = []
        for col in df.columns:
            dtype = str(df[col].dtype)
            col_type = 'numeric' if ('int' in dtype or 'float' in dtype) else ('date' if 'datetime' in dtype else 'string')
            col_info.append({'name': str(col), 'type': col_type, 'dtype': dtype, 'missing': int(df[col].isna().sum())})
        dataset.set_column_info(col_info)
        dataset.save()

        return JsonResponse({'success': True, 'rows': len(df), 'columns': len(df.columns), 'column_info': col_info})

    except Exception as e:
        import traceback
        return JsonResponse({'error': str(e), 'traceback': traceback.format_exc()}, status=500)



# ──────────────────────────────────────────────────────────────────────────────
# ANALYSIS API
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def run_analysis(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    dataset_id = body.get('dataset_id')
    analysis_type = body.get('analysis_type')
    params = body.get('params', {})

    if not dataset_id or not analysis_type:
        return JsonResponse({'error': 'dataset_id and analysis_type are required'}, status=400)

    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)

    try:
        df = load_df(dataset)
        engine = AnalysisEngine(df)
        output_html = ''

        if analysis_type == 'descriptive':
            variables = params.get('variables', [])
            output_html = engine.descriptive(variables)

        elif analysis_type == 'frequencies':
            variables = params.get('variables', [])
            output_html = engine.frequencies(variables)

        elif analysis_type == 'crosstab':
            output_html = engine.crosstab(params['row_var'], params['col_var'])

        elif analysis_type == 'ttest_independent':
            output_html = engine.ttest_independent(
                params['dep_var'], params['group_var'],
                params['group1'], params['group2']
            )

        elif analysis_type == 'ttest_onesample':
            output_html = engine.ttest_onesample(
                params['variables'], float(params.get('test_value', 0))
            )

        elif analysis_type == 'ttest_paired':
            pairs = [(p['var1'], p['var2']) for p in params.get('pairs', [])]
            output_html = engine.ttest_paired(pairs)

        elif analysis_type == 'anova_oneway':
            output_html = engine.anova_oneway(params['dep_var'], params['factor_var'])

        elif analysis_type == 'anova_twoway':
            output_html = engine.anova_twoway(params['dep_var'], params['factor1'], params['factor2'])

        elif analysis_type == 'manova':
            output_html = engine.manova(params['dep_vars'], params['factor_var'])

        elif analysis_type == 'correlation_pearson':
            output_html = engine.correlation_pearson(params['variables'])

        elif analysis_type == 'correlation_spearman':
            output_html = engine.correlation_spearman(params['variables'])

        elif analysis_type == 'regression_linear':
            output_html = engine.regression_linear(params['dep_var'], params['indep_vars'])

        elif analysis_type == 'regression_logistic':
            output_html = engine.regression_logistic(params['dep_var'], params['indep_vars'])

        elif analysis_type == 'normality':
            output_html = engine.normality(params['variables'])

        elif analysis_type == 'chi_square':
            output_html = engine.chi_square(params['row_var'], params['col_var'])

        elif analysis_type == 'mann_whitney':
            output_html = engine.mann_whitney(
                params['dep_var'], params['group_var'],
                params['group1'], params['group2']
            )

        elif analysis_type == 'wilcoxon':
            output_html = engine.wilcoxon(params['var1'], params['var2'])

        elif analysis_type == 'kruskal_wallis':
            output_html = engine.kruskal_wallis(params['dep_var'], params['group_var'])

        elif analysis_type == 'factor':
            output_html = engine.factor_analysis(
                params['variables'], int(params.get('n_factors', 0)) or None
            )

        elif analysis_type == 'cluster':
            output_html = engine.cluster_analysis(
                params['variables'], int(params.get('n_clusters', 3))
            )

        elif analysis_type == 'reliability':
            output_html = engine.reliability_alpha(params['variables'])

        else:
            return JsonResponse({'error': f'Unknown analysis type: {analysis_type}'}, status=400)

        # Save result
        result = AnalysisResult.objects.create(
            dataset=dataset,
            analysis_type=analysis_type,
            parameters=json.dumps(params),
            output_html=output_html,
        )

        return JsonResponse({
            'success': True,
            'result_id': result.id,
            'output_html': output_html,
            'analysis_type': analysis_type,
        })

    except Exception as e:
        import traceback
        return JsonResponse({'error': str(e), 'traceback': traceback.format_exc()}, status=500)


@login_required
def generate_chart(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    dataset_id = body.get('dataset_id')
    chart_type = body.get('chart_type')
    params = body.get('params', {})

    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)

    try:
        df = load_df(dataset)
        gen = ChartGenerator(df)
        chart_json = {}

        if chart_type == 'bar':
            chart_json = gen.bar_chart(params.get('x_var'), params.get('y_var'), params.get('color_var'), params)
        elif chart_type == 'histogram':
            chart_json = gen.histogram(params.get('x_var'), params.get('bins'), params.get('color_var'), params)
        elif chart_type == 'scatter':
            chart_json = gen.scatter_plot(params.get('x_var'), params.get('y_var'), params.get('color_var'), params.get('size_var'), params)
        elif chart_type == 'box':
            chart_json = gen.box_plot(params.get('y_var'), params.get('x_var'), params)
        elif chart_type == 'line':
            y_vars = params.get('y_vars', params.get('y_var', []))
            if isinstance(y_vars, str):
                y_vars = [y_vars]
            chart_json = gen.line_chart(params.get('x_var'), y_vars, params)
        elif chart_type == 'pie':
            chart_json = gen.pie_chart(params.get('names_var'), params.get('values_var'), params)
        else:
            return JsonResponse({'error': 'Unknown chart type'}, status=400)

        # Save chart
        ChartResult.objects.create(
            dataset=dataset,
            chart_type=chart_type,
            parameters=json.dumps(params),
            chart_json=json.dumps(chart_json),
        )

        return JsonResponse({'success': True, 'chart_json': chart_json})

    except Exception as e:
        import traceback
        return JsonResponse({'error': str(e), 'traceback': traceback.format_exc()}, status=500)


# ──────────────────────────────────────────────────────────────────────────────
# EXPORT
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def export_output(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        body = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    fmt = body.get('format', 'pdf')
    output_html = body.get('output_html', '')

    if fmt == 'pdf':
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import cm
            import re

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4,
                                    leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            story = []

            # Strip HTML tags
            clean = re.sub('<[^<]+?>', '', output_html)
            for para in clean.split('\n'):
                para = para.strip()
                if para:
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 6))

            doc.build(story)
            buf.seek(0)
            response = HttpResponse(buf.read(), content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="spss_output.pdf"'
            return response
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    elif fmt == 'docx':
        try:
            import docx
            import re
            buf = io.BytesIO()
            doc = docx.Document()
            doc.add_heading('SPSS Online - Output Results', 0)
            clean = re.sub('<[^<]+?>', '', output_html)
            for para in clean.split('\n'):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
            doc.save(buf)
            buf.seek(0)
            response = HttpResponse(buf.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="spss_output.docx"'
            return response
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Format tidak didukung'}, status=400)


# ──────────────────────────────────────────────────────────────────────────────
# DATA EDITOR OPERATIONS
# ──────────────────────────────────────────────────────────────────────────────

@login_required
def add_row(request):
    """Add empty row(s) to dataset."""
    return JsonResponse({'success': True, 'message': 'Row added (client-side)'})


@login_required
def get_analysis_history(request, dataset_id):
    dataset = get_object_or_404(Dataset, id=dataset_id, user=request.user)
    analyses = AnalysisResult.objects.filter(dataset=dataset).order_by('-created_at')[:20]
    result = [{
        'id': a.id,
        'type': a.analysis_type,
        'type_label': dict(AnalysisResult.ANALYSIS_TYPES).get(a.analysis_type, a.analysis_type),
        'created_at': a.created_at.strftime('%H:%M:%S'),
        'output_html': a.output_html,
    } for a in analyses]
    return JsonResponse({'history': result})
